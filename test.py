import os
import docx


def edit_application(procedure_number, quantity):
    path_to_doc = os.getcwd().replace('\\', '/') + '/docs-2/Заявка Митра.docx'
    doc = docx.Document(path_to_doc)
    price = str(int(quantity.split('.')[0]) * 0.01)

    doc.paragraphs[2].runs[2].text = procedure_number
    doc.paragraphs[9].runs[1].text = procedure_number
    doc.paragraphs[9].runs[3].text = price
    doc.paragraphs[11].runs[1].text = quantity.split('.')[0]

    doc.save(os.getcwd().replace('\\', '/') + '/docs-2/Заявка Митра-test.docx')

edit_application('123', '3')