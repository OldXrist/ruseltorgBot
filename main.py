import docx
from selenium import webdriver
from selenium.webdriver import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from twocaptcha import TwoCaptcha
import multiprocessing
import time
import os

config = {
    'server': 'rucaptcha.com',
    'apiKey': '270776e78e5170b3d26e8cebbb713152',
}
solver = TwoCaptcha(**config)


def edit_application(procedure_number, quantity):
    path_to_doc = os.getcwd().replace('\\', '/') + '/docs-2/Заявка Право-Торг.docx'
    doc = docx.Document(path_to_doc)
    price = str(int(quantity.split('.')[0]) * 0.01)

    doc.paragraphs[2].runs[2].text = procedure_number
    doc.paragraphs[9].runs[1].text = procedure_number
    doc.paragraphs[9].runs[6].text = price
    doc.paragraphs[11].runs[1].text = quantity.split('.')[0]

    doc.save(path_to_doc)


def get_captcha_sync(driver, main_w, captcha_w):
    driver.switch_to.window(captcha_w)
    driver.refresh()
    time.sleep(1)
    with open('captcha.png', 'wb') as file:
        file.write(chromeDriver.find_elements(By.CSS_SELECTOR, "img")[0].screenshot_as_png)
        file.close()
    driver.switch_to.window(main_w)
    result = solver.normal(os.getcwd().replace('\\', '/') + '/captcha.png', language='en', minLength=6, maxLength=6)['code']
    with open('captcha.txt', 'w') as file:
        file.write(result)
        file.close()


def fill_captcha(driver, main_w, captcha_w):
    with open('captcha.txt', 'r') as file:
        result = file.read()
        file.close()
    for i in range(10):
        driver.find_elements(By.CSS_SELECTOR, ".x-form-text.x-form-field")[len(driver.find_elements(By.CSS_SELECTOR, ".x-form-text.x-form-field")) - 1].send_keys(Keys.BACKSPACE)
    driver.find_elements(By.CSS_SELECTOR, ".x-form-text.x-form-field")[len(driver.find_elements(By.CSS_SELECTOR, ".x-form-text.x-form-field")) - 1].send_keys(result)
    buttons = driver.find_elements(By.CSS_SELECTOR, ".x-btn-text")
    for button in buttons:
        if button.get_attribute("innerText") == "Отправить":
            button.click()
    time.sleep(1)
    errors = driver.find_elements(By.CSS_SELECTOR, ".x-window-header")
    for error in errors:
        if error.find_elements(By.CSS_SELECTOR, ".x-window-header-text")[0].get_attribute("innerText") == "Ошибка":
            error.find_elements(By.CSS_SELECTOR, ".x-tool-close")[0].click()
            get_captcha_sync(driver, main_w, captcha_w)


def save_new_captcha():
    result = solver.normal(os.getcwd().replace('\\', '/') + '/captcha.png', language='en', minLength=6, maxLength=6)['code']
    with open('captcha.txt', 'w') as file:
        file.write(result)
        file.close()


def read_new_captcha(driver, main_w, captcha_w):
    driver.switch_to.window(captcha_w)
    driver.refresh()
    time.sleep(1)
    with open('captcha.png', 'wb') as file:
        file.write(chromeDriver.find_elements(By.CSS_SELECTOR, "img")[0].screenshot_as_png)
        file.close()
    driver.switch_to.window(main_w)
    p = multiprocessing.Process(target=save_new_captcha)
    p.start()


def main(driver, main_w, captcha_w):

    get_captcha_sync(driver, main_w, captcha_w)
    while True:
        if driver.find_elements(By.CSS_SELECTOR, "#captcha_img"):
            fill_captcha(driver, main_w, captcha_w)
            time.sleep(1)
        else:
            break

    read_new_captcha(driver, main_w, captcha_w)

    if driver.find_element(By.ID, "ext-gen232"):
        driver.find_element(By.ID, "ext-gen232").click()
    time.sleep(2)

    if driver.find_element(By.ID, "ext-gen524"):
        driver.execute_script("document.getElementById('ext-gen524').click()")
    time.sleep(2)

    if driver.find_element(By.ID, "customers_name"):
        driver.find_element(By.ID, "customers_name").send_keys("7708701670")
    time.sleep(2)

    if driver.find_element(By.ID, "ext-gen300"):
        driver.execute_script("document.getElementById('ext-gen300').click()")

    while True:
        while True:
            if driver.find_elements(By.CSS_SELECTOR, "#captcha_img"):
                read_new_captcha(driver, main_w, captcha_w)
                fill_captcha(driver, main_w, captcha_w)
                time.sleep(1)
            else:
                break

        if driver.find_elements(By.CSS_SELECTOR, ".x-grid3-cell-inner.x-grid3-col-1"):
            lotNum = driver.find_elements(By.CSS_SELECTOR, ".x-grid3-cell-inner.x-grid3-col-1")[0].get_attribute("innerText")
            with open('lot_numbers.txt', 'r') as r:
                if lotNum in r.read():
                    driver.execute_script("document.getElementsByClassName('x-grid3-row')[0].remove()")
                else:
                    with open('lot_numbers.txt', 'a') as f:
                        f.write(lotNum)
                        if driver.find_elements(By.CSS_SELECTOR, ".x-grid3-cell-inner.x-grid3-col-14")[0]:
                            linkContainer = driver.find_elements(By.CSS_SELECTOR, ".x-grid3-cell-inner.x-grid3-col-14")[0]
                            linkContainer.find_elements(By.TAG_NAME, "a")[0].click()
                        break
        else:
            driver.execute_script("document.getElementById('ext-gen300').click()")
            time.sleep(3)

    time.sleep(5)

    quantity = driver.find_elements(By.CSS_SELECTOR, ".x-grid3-col-4")[0].get_attribute('innerHTML')
    tr_list = driver.find_elements(By.CSS_SELECTOR, "tr")
    for i in range(len(tr_list)):
        if 'Закупка №:' in tr_list[i].get_attribute('innerText'):
            trade_number = tr_list[i].find_elements(By.CSS_SELECTOR, "a")[0].get_attribute('innerText')
            edit_application(trade_number, quantity)
            break

    trade_id = driver.current_url.split('/')[len(driver.current_url.split('/')) - 1]

    driver.get("https://etp.roseltorg.ru/supplier/auction/apply/auction_id/" + trade_id)

    if driver.find_elements(By.CSS_SELECTOR, "#origin_preferences_not_provided"):
        driver.execute_script("document.getElementById('origin_preferences_not_provided').click()")

    if driver.find_elements(By.CSS_SELECTOR, "#declaration_check_text"):
        driver.execute_script("document.getElementById('declaration_check_text').click()")

    if driver.find_elements(By.CSS_SELECTOR, "#acct_ras"):
        driver.find_element(By.ID, "acct_ras").send_keys("40702810026070002060")

    if driver.find_elements(By.CSS_SELECTOR, "#acct_kor"):
        driver.find_element(By.ID, "acct_kor").send_keys("30101810500000000207")

    if driver.find_elements(By.CSS_SELECTOR, "#bik"):
        driver.find_element(By.ID, "bik").send_keys("046015207")

        time.sleep(2)

    if driver.find_elements(By.CSS_SELECTOR, ".search-item"):
        driver.find_elements(By.CSS_SELECTOR, ".search-item")[0].click()
    else:
        if driver.find_elements(By.CSS_SELECTOR, "#bank"):
            driver.find_element(By.ID, "bank").send_keys("ФИЛИАЛ \"РОСТОВСКИЙ\" АО \"АЛЬФА-БАНК\"")

        if driver.find_elements(By.CSS_SELECTOR, "#bank_addr"):
            driver.find_element(By.ID, "bank_addr").send_keys("Ростов-на-Дону, ул Красноармейская, 206")

    if driver.find_element(By.CSS_SELECTOR, "#files_maxsum_panel_i"):
        permissionField = driver.find_element(By.CSS_SELECTOR, "#files_maxsum_panel_i")
        if permissionField.find_elements(By.CSS_SELECTOR, ".x-form-file"):
            file = permissionField.find_elements(By.CSS_SELECTOR, ".x-form-file")[0]
            file.send_keys(os.getcwd().replace('\\', '/') + "/docs-1/Об одобрении крупнои сделки.pdf")

    if driver.find_elements(By.CSS_SELECTOR, "#contract_price_offer"):
        driver.find_element(By.ID, "contract_price_offer").send_keys(int(quantity.split('.')[0]) * 0.01)

    otherFilesSection = driver.find_element(By.ID, "files_type4_panel")
    docAddButton = driver.find_element(By.ID, "files_type4_btn")

    directory = os.getcwd().replace('\\', '/') + "/docs-2/"

    for filename in os.listdir(directory):
        docAddButton.click()
        otherFiles = otherFilesSection.find_elements(By.CSS_SELECTOR, ".x-form-file")

        otherFiles[0].send_keys(directory + filename)

    buttons = driver.find_elements(By.CSS_SELECTOR, "button")

    for i in range(len(buttons)):
        if buttons[i].get_attribute('innerHTML') == "Подписать и направить заявку":
            # buttons[i].click()
            time.sleep(500)


if __name__ == '__main__':
    options = Options()
    cwd = os.getcwd().replace('\\', '/') + "/UserData"
    options.add_argument("--user-data-dir=" + cwd)
    options.page_load_strategy = 'normal'
    chromeDriver = webdriver.Chrome(options=options)
    mainWindow = chromeDriver.window_handles[0]
    chromeDriver.get("https://etp.roseltorg.ru/supplier/auction/index")
    #  time.sleep(5000)  # для входа в аккаунт
    chromeDriver.execute_script("window.open('about:blank', 'captchaWindow');")
    time.sleep(1)
    captchaWindow = chromeDriver.window_handles[1]
    chromeDriver.switch_to.window(captchaWindow)
    chromeDriver.get("https://etp.roseltorg.ru/data/auctionlist2.php?captcha=1")
    chromeDriver.switch_to.window(mainWindow)
    main(chromeDriver, mainWindow, captchaWindow)
